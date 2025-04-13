import os
import chess
import chess.pgn
import chess.svg
from docx import Document
from docx.shared import Inches
from io import BytesIO
from cairosvg import svg2png
from natsort import natsorted  # <-- Added for natural sorting

def create_puzzle_book(pgn_folder, output_doc="chess_puzzles.docx"):
    doc = Document()
    solutions = []
    doc.add_heading("Chess Puzzle Book", 0)

    # Natural sorting for files (e.g. puzzle1.pgn, puzzle2.pgn... puzzle10.pgn)
    for idx, pgn_file in enumerate(natsorted(os.listdir(pgn_folder))):  #Fixed sorting
 
        if not pgn_file.lower().endswith('.pgn'):
            continue

        pgn_path = os.path.join(pgn_folder, pgn_file)
        try:
            with open(pgn_path) as f:
                game = chess.pgn.read_game(f)
                if not game:
                    print(f"Skipping {pgn_file}: Empty or invalid PGN")
                    continue

                board = game.board()
                mainline_moves = list(game.mainline_moves())

                if not mainline_moves:
                    print(f"Skipping {pgn_file}: No moves found")
                    continue

                first_move = mainline_moves[0]
                board.push(first_move)
                turn = "White" if board.turn == chess.WHITE else "Black"

                svg = chess.svg.board(board=board)
                png_data = svg2png(bytestring=svg.encode('utf-8'))
                image_stream = BytesIO(png_data)

                doc.add_heading(f"Puzzle {idx+1} ({turn} to move)", level=2)
                doc.add_picture(image_stream, width=Inches(4.0))
                doc.add_page_break()

                solutions.append(game.accept(chess.pgn.StringExporter(headers=True)))

        except Exception as e:
            print(f"Error processing {pgn_file}: {e}")
            continue

    if solutions:
        doc.add_heading("Solutions", 1)
        for idx, pgn in enumerate(solutions):
            doc.add_heading(f"Solution {idx+1}", level=2)
            doc.add_paragraph(pgn)
            doc.add_page_break()

    doc.save(output_doc)
    print(f"Created {output_doc} with {len(solutions)} puzzles!")

if __name__ == "__main__":
    pgn_folder = "C:/Users/Student/Desktop/Python programs/Chess book/pgn_folder"
    create_puzzle_book(pgn_folder)
