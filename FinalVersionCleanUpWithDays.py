import os
import re
import chess
import chess.pgn
import chess.svg
from docx import Document
from docx.shared import Inches
from io import BytesIO, StringIO
from cairosvg import svg2png
from natsort import natsorted

def clean_pgn_content(content):
    """Remove comments, variations, and unwanted headers from PGN content."""
    def remove_nested(text, pattern):
        while re.search(pattern, text):
            text = re.sub(pattern, '', text)
        return text

    content = remove_nested(content, r'\([^()]*\)')
    content = remove_nested(content, r'\{[^{}]*\}')

    unwanted_headers = ["Link", "Site", "Event", "Date", "Round"]
    for header in unwanted_headers:
        content = re.sub(rf'\[{header} ".*?"\]\s*', '', content)

    return content

def create_puzzle_book(pgn_folder, output_doc="chess_puzzles.docx"):
    doc = Document()
    solutions = []
    doc.add_heading("Chess Puzzle Book", 0)

    puzzle_count = 0

    for idx, pgn_file in enumerate(natsorted(os.listdir(pgn_folder))):
        if not pgn_file.lower().endswith('.pgn'):
            continue

        pgn_path = os.path.join(pgn_folder, pgn_file)
        try:
            with open(pgn_path, "r", encoding="utf-8") as f:
                raw_content = f.read()

            cleaned_content = clean_pgn_content(raw_content)
            game = chess.pgn.read_game(StringIO(cleaned_content))

            if not game:
                print(f"Skipping {pgn_file}: Empty or invalid PGN")
                continue

            board = game.board()
            mainline_moves = list(game.mainline_moves())

            if not mainline_moves:
                print(f"Skipping {pgn_file}: No moves found")
                continue

            puzzle_count += 1

            # Insert a "Day N" heading every 10 puzzles
            if (puzzle_count - 1) % 10 == 0:
                day_number = (puzzle_count - 1) // 10 + 1
                doc.add_heading(f"Day {day_number}", level=2)

            first_move = mainline_moves[0]
            board.push(first_move)
            turn = "White" if board.turn == chess.WHITE else "Black"

            svg = chess.svg.board(board=board)
            png_data = svg2png(bytestring=svg.encode('utf-8'))
            image_stream = BytesIO(png_data)

            doc.add_heading(f"Puzzle {puzzle_count} ({turn} to move)", level=3)
            doc.add_picture(image_stream, width=Inches(3.0))

            # Add a single line breaks instead of page break
            doc.add_paragraph()
           
            exporter = chess.pgn.StringExporter(headers=True)
            solutions.append(game.accept(exporter))

        except Exception as e:
            print(f"Error processing {pgn_file}: {e}")
            continue

    if solutions:
        doc.add_heading("Solutions", 1)
        for idx, pgn in enumerate(solutions):
            doc.add_heading(f"Solution {idx+1}", level=2)
            doc.add_paragraph(pgn)
            doc.add_page_break()

    doc.save(output_doc)
    print(f"Created {output_doc} with {len(solutions)} puzzles!")

if __name__ == "__main__":
    pgn_folder = "C:/Users/Student/Desktop/Python programs/Chess book/pgn_folder"
    create_puzzle_book(pgn_folder)
